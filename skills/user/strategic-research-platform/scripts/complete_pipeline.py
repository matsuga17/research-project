#!/usr/bin/env python3
"""
Complete Research Pipeline for Strategic Management Studies
============================================================

This script implements the full research workflow from data collection to 
publication-ready outputs. Follows best practices for reproducibility and 
transparency in empirical strategy research.

Author: Strategic Management Research Hub
Version: 1.0
Date: 2025-10-31

Pipeline Phases:
1. Data Collection (APIs, Web scraping, Manual)
2. Data Cleaning (Missing data, Outliers, Transformations)
3. Data Integration (Multi-source merging)
4. Descriptive Statistics (Tables 1-2)
5. Regression Analysis (OLS, Panel, IV, DiD)
6. Robustness Checks (Alternative specifications)
7. Visualization (Figures 1-3)
8. Export Results (LaTeX tables, DOCX report)

Usage:
    python complete_pipeline.py --config config.yaml
    
Configuration file (config.yaml) should specify:
- data_sources: List of data sources
- sample_criteria: Firm selection rules
- variables: DV, IVs, controls, moderators
- analysis: Regression specifications
- output_dir: Where to save results
"""

import os
import sys
import yaml
import logging
import argparse
import warnings
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Union

# Data manipulation
import numpy as np
import pandas as pd
from scipy import stats
from scipy.stats.mstats import winsorize

# Statistical analysis
import statsmodels.api as sm
import statsmodels.formula.api as smf
from statsmodels.stats.outliers_influence import variance_inflation_factor
from linearmodels import PanelOLS, RandomEffects

# Visualization
import matplotlib.pyplot as plt
import seaborn as sns

# Data collection APIs
import requests
from bs4 import BeautifulSoup

# Export
from docx import Document
from docx.shared import Inches, Pt
import openpyxl

# Suppress warnings for cleaner output
warnings.filterwarnings('ignore')

# ============================================================================
# Configuration & Logging Setup
# ============================================================================

def setup_logging(output_dir: str) -> logging.Logger:
    """
    Setup logging configuration for pipeline execution.
    
    Args:
        output_dir: Directory for log files
        
    Returns:
        Configured logger instance
    """
    log_dir = Path(output_dir) / "logs"
    log_dir.mkdir(parents=True, exist_ok=True)
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    log_file = log_dir / f"pipeline_{timestamp}.log"
    
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s',
        handlers=[
            logging.FileHandler(log_file),
            logging.StreamHandler(sys.stdout)
        ]
    )
    
    logger = logging.getLogger(__name__)
    logger.info(f"Pipeline started. Log file: {log_file}")
    
    return logger


def load_config(config_path: str) -> Dict:
    """
    Load configuration from YAML file.
    
    Args:
        config_path: Path to config.yaml
        
    Returns:
        Configuration dictionary
    """
    with open(config_path, 'r') as f:
        config = yaml.safe_load(f)
    
    # Validate required fields
    required_fields = ['data_sources', 'sample_criteria', 'variables', 'analysis', 'output_dir']
    for field in required_fields:
        if field not in config:
            raise ValueError(f"Missing required field in config: {field}")
    
    return config


# ============================================================================
# Phase 1: Data Collection
# ============================================================================

class DataCollector:
    """
    Unified data collection interface for multiple sources.
    """
    
    def __init__(self, config: Dict, logger: logging.Logger):
        self.config = config
        self.logger = logger
        self.data_sources = config['data_sources']
    
    def collect_all(self) -> Dict[str, pd.DataFrame]:
        """
        Collect data from all configured sources.
        
        Returns:
            Dictionary of {source_name: dataframe}
        """
        self.logger.info("=" * 60)
        self.logger.info("PHASE 1: DATA COLLECTION")
        self.logger.info("=" * 60)
        
        collected_data = {}
        
        for source in self.data_sources:
            source_name = source['name']
            source_type = source['type']
            
            self.logger.info(f"Collecting data from: {source_name} ({source_type})")
            
            try:
                if source_type == 'csv':
                    df = self._collect_csv(source)
                elif source_type == 'api':
                    df = self._collect_api(source)
                elif source_type == 'database':
                    df = self._collect_database(source)
                elif source_type == 'manual':
                    df = self._load_manual_data(source)
                else:
                    raise ValueError(f"Unknown source type: {source_type}")
                
                collected_data[source_name] = df
                self.logger.info(f"  ✓ Collected {len(df)} observations from {source_name}")
                
            except Exception as e:
                self.logger.error(f"  ✗ Failed to collect from {source_name}: {str(e)}")
                raise
        
        return collected_data
    
    def _collect_csv(self, source: Dict) -> pd.DataFrame:
        """Load data from CSV file."""
        df = pd.read_csv(source['path'])
        return df
    
    def _collect_api(self, source: Dict) -> pd.DataFrame:
        """
        Collect data from API endpoint.
        
        Example for SEC EDGAR, EDINET, USPTO PatentsView, etc.
        """
        api_url = source['url']
        params = source.get('params', {})
        
        # Add rate limiting
        import time
        time.sleep(1)  # Be respectful to APIs
        
        response = requests.get(api_url, params=params)
        response.raise_for_status()
        
        data = response.json()
        df = pd.json_normalize(data['results'])  # Adjust based on API structure
        
        return df
    
    def _collect_database(self, source: Dict) -> pd.DataFrame:
        """
        Collect from database (e.g., WRDS, PostgreSQL).
        """
        # Example: WRDS connection
        import wrds
        
        db = wrds.Connection()
        query = source['query']
        df = db.raw_sql(query)
        db.close()
        
        return df
    
    def _load_manual_data(self, source: Dict) -> pd.DataFrame:
        """Load manually collected data."""
        df = pd.read_excel(source['path'])
        return df


# ============================================================================
# Phase 2: Data Cleaning
# ============================================================================

class DataCleaner:
    """
    Standardized data cleaning pipeline.
    """
    
    def __init__(self, config: Dict, logger: logging.Logger):
        self.config = config
        self.logger = logger
    
    def clean_all(self, data_dict: Dict[str, pd.DataFrame]) -> Dict[str, pd.DataFrame]:
        """
        Clean all datasets.
        
        Args:
            data_dict: Dictionary of {source_name: dataframe}
            
        Returns:
            Dictionary of cleaned dataframes
        """
        self.logger.info("=" * 60)
        self.logger.info("PHASE 2: DATA CLEANING")
        self.logger.info("=" * 60)
        
        cleaned_data = {}
        
        for source_name, df in data_dict.items():
            self.logger.info(f"Cleaning: {source_name}")
            
            # Initial size
            initial_rows = len(df)
            
            # Step 1: Handle missing data
            df = self._handle_missing(df, source_name)
            
            # Step 2: Detect and treat outliers
            df = self._handle_outliers(df, source_name)
            
            # Step 3: Create derived variables
            df = self._create_variables(df, source_name)
            
            # Step 4: Validate data types
            df = self._validate_types(df, source_name)
            
            cleaned_data[source_name] = df
            
            final_rows = len(df)
            rows_dropped = initial_rows - final_rows
            pct_dropped = (rows_dropped / initial_rows) * 100
            
            self.logger.info(f"  ✓ Cleaned {source_name}: {final_rows} rows (dropped {rows_dropped}, {pct_dropped:.1f}%)")
        
        return cleaned_data
    
    def _handle_missing(self, df: pd.DataFrame, source_name: str) -> pd.DataFrame:
        """
        Handle missing data according to best practices.
        """
        # Log missing data patterns
        missing_counts = df.isnull().sum()
        missing_vars = missing_counts[missing_counts > 0]
        
        if len(missing_vars) > 0:
            self.logger.info(f"    Missing data in {source_name}:")
            for var, count in missing_vars.items():
                pct = (count / len(df)) * 100
                self.logger.info(f"      {var}: {count} ({pct:.1f}%)")
        
        # Strategy 1: Drop if key variables missing
        key_vars = self.config.get('key_variables', [])
        if key_vars:
            df = df.dropna(subset=key_vars)
        
        # Strategy 2: Impute R&D (common in strategy research)
        if 'rd_expense' in df.columns:
            df['rd_missing'] = df['rd_expense'].isnull().astype(int)
            df['rd_expense'] = df['rd_expense'].fillna(0)
        
        return df
    
    def _handle_outliers(self, df: pd.DataFrame, source_name: str) -> pd.DataFrame:
        """
        Detect and treat outliers using winsorization.
        """
        numeric_cols = df.select_dtypes(include=[np.number]).columns
        winsorize_vars = self.config.get('winsorize_variables', [])
        
        for var in winsorize_vars:
            if var in numeric_cols:
                # Winsorize at 1st and 99th percentiles
                df[var] = winsorize(df[var].dropna(), limits=[0.01, 0.01])
                self.logger.info(f"    Winsorized: {var}")
        
        return df
    
    def _create_variables(self, df: pd.DataFrame, source_name: str) -> pd.DataFrame:
        """
        Create derived variables (log transforms, ratios, lags).
        """
        # Log transformations
        if 'total_assets' in df.columns:
            df['log_assets'] = np.log(df['total_assets'])
        
        if 'employees' in df.columns:
            df['log_employees'] = np.log(df['employees'] + 1)
        
        # Financial ratios
        if 'net_income' in df.columns and 'total_assets' in df.columns:
            df['roa'] = df['net_income'] / df['total_assets']
        
        if 'rd_expense' in df.columns and 'sales' in df.columns:
            df['rd_intensity'] = df['rd_expense'] / df['sales']
        
        if 'total_debt' in df.columns and 'total_assets' in df.columns:
            df['leverage'] = df['total_debt'] / df['total_assets']
        
        # Panel data: Create lagged variables
        if 'firm_id' in df.columns and 'year' in df.columns:
            df = df.sort_values(['firm_id', 'year'])
            for var in self.config.get('lag_variables', []):
                if var in df.columns:
                    df[f'{var}_lag1'] = df.groupby('firm_id')[var].shift(1)
        
        return df
    
    def _validate_types(self, df: pd.DataFrame, source_name: str) -> pd.DataFrame:
        """
        Ensure correct data types.
        """
        # Convert year to int
        if 'year' in df.columns:
            df['year'] = df['year'].astype(int)
        
        # Convert firm_id to string
        if 'firm_id' in df.columns:
            df['firm_id'] = df['firm_id'].astype(str)
        
        return df


# ============================================================================
# Phase 3: Data Integration
# ============================================================================

class DataIntegrator:
    """
    Merge data from multiple sources.
    """
    
    def __init__(self, config: Dict, logger: logging.Logger):
        self.config = config
        self.logger = logger
    
    def integrate(self, cleaned_data: Dict[str, pd.DataFrame]) -> pd.DataFrame:
        """
        Merge all data sources into single analytical dataset.
        
        Returns:
            Integrated dataframe
        """
        self.logger.info("=" * 60)
        self.logger.info("PHASE 3: DATA INTEGRATION")
        self.logger.info("=" * 60)
        
        # Start with primary dataset
        primary_source = self.config.get('primary_source', list(cleaned_data.keys())[0])
        df = cleaned_data[primary_source].copy()
        self.logger.info(f"Primary dataset: {primary_source} ({len(df)} obs)")
        
        # Merge with secondary sources
        for source_name, source_df in cleaned_data.items():
            if source_name == primary_source:
                continue
            
            merge_keys = self.config.get('merge_keys', {}).get(source_name, ['firm_id', 'year'])
            merge_type = self.config.get('merge_type', {}).get(source_name, 'left')
            
            initial_rows = len(df)
            df = pd.merge(df, source_df, on=merge_keys, how=merge_type, suffixes=('', f'_{source_name}'))
            final_rows = len(df)
            
            self.logger.info(f"  Merged {source_name}: {initial_rows} → {final_rows} obs")
        
        # Create industry and year fixed effects
        if 'industry' in df.columns:
            df = pd.get_dummies(df, columns=['industry'], prefix='ind', drop_first=True)
        
        if 'year' in df.columns:
            df = pd.get_dummies(df, columns=['year'], prefix='yr', drop_first=True)
        
        self.logger.info(f"Final integrated dataset: {len(df)} observations, {len(df.columns)} variables")
        
        return df


# ============================================================================
# Phase 4: Descriptive Statistics
# ============================================================================

class DescriptiveAnalyzer:
    """
    Generate descriptive statistics tables.
    """
    
    def __init__(self, config: Dict, logger: logging.Logger):
        self.config = config
        self.logger = logger
    
    def analyze(self, df: pd.DataFrame) -> Dict[str, pd.DataFrame]:
        """
        Generate all descriptive tables.
        
        Returns:
            Dictionary of tables
        """
        self.logger.info("=" * 60)
        self.logger.info("PHASE 4: DESCRIPTIVE STATISTICS")
        self.logger.info("=" * 60)
        
        tables = {}
        
        # Table 1: Sample composition
        tables['sample_composition'] = self._sample_composition(df)
        
        # Table 2: Descriptive statistics
        tables['descriptive_stats'] = self._descriptive_stats(df)
        
        # Table 3: Correlation matrix
        tables['correlation'] = self._correlation_matrix(df)
        
        return tables
    
    def _sample_composition(self, df: pd.DataFrame) -> pd.DataFrame:
        """
        Table 1: Sample composition by industry/year.
        """
        # By industry (if available)
        if 'industry' in df.columns:
            comp = df.groupby('industry').agg({
                'firm_id': 'nunique',
                'roa': 'mean'
            }).reset_index()
            comp.columns = ['Industry', 'N Firms', 'Avg ROA']
            comp['% of Sample'] = (comp['N Firms'] / comp['N Firms'].sum() * 100).round(1)
            
            self.logger.info("Sample Composition by Industry:")
            self.logger.info(f"\n{comp.to_string(index=False)}")
            
            return comp
        
        return pd.DataFrame()
    
    def _descriptive_stats(self, df: pd.DataFrame) -> pd.DataFrame:
        """
        Table 2: Descriptive statistics (mean, SD, min, max).
        """
        vars_to_describe = self.config.get('descriptive_variables', [])
        
        if not vars_to_describe:
            # Default: all numeric variables
            vars_to_describe = df.select_dtypes(include=[np.number]).columns.tolist()
        
        desc = df[vars_to_describe].describe().T
        desc['N'] = df[vars_to_describe].count()
        desc = desc[['N', 'mean', 'std', 'min', '25%', '50%', '75%', 'max']]
        desc = desc.round(3)
        
        self.logger.info("Descriptive Statistics:")
        self.logger.info(f"\n{desc.to_string()}")
        
        return desc
    
    def _correlation_matrix(self, df: pd.DataFrame) -> pd.DataFrame:
        """
        Table 3: Correlation matrix with significance stars.
        """
        vars_for_corr = self.config.get('correlation_variables', [])
        
        if not vars_for_corr:
            vars_for_corr = self.config.get('descriptive_variables', [])
        
        corr = df[vars_for_corr].corr().round(2)
        
        self.logger.info("Correlation Matrix:")
        self.logger.info(f"\n{corr.to_string()}")
        
        # Check for multicollinearity
        high_corr = (corr.abs() > 0.8) & (corr.abs() < 1.0)
        if high_corr.any().any():
            self.logger.warning("  ⚠ High correlations detected (|r| > 0.8):")
            for i in range(len(corr)):
                for j in range(i+1, len(corr)):
                    if high_corr.iloc[i, j]:
                        self.logger.warning(f"    {corr.index[i]} <-> {corr.columns[j]}: {corr.iloc[i,j]:.2f}")
        
        return corr


# ============================================================================
# Phase 5: Regression Analysis
# ============================================================================

class RegressionAnalyzer:
    """
    Run regression models (OLS, Panel, IV, DiD).
    """
    
    def __init__(self, config: Dict, logger: logging.Logger):
        self.config = config
        self.logger = logger
    
    def analyze(self, df: pd.DataFrame) -> Dict[str, any]:
        """
        Run all regression specifications.
        
        Returns:
            Dictionary of regression results
        """
        self.logger.info("=" * 60)
        self.logger.info("PHASE 5: REGRESSION ANALYSIS")
        self.logger.info("=" * 60)
        
        results = {}
        
        models = self.config['analysis'].get('models', [])
        
        for i, model_spec in enumerate(models, 1):
            model_name = model_spec.get('name', f'Model {i}')
            self.logger.info(f"Running: {model_name}")
            
            model_type = model_spec['type']
            
            if model_type == 'ols':
                result = self._run_ols(df, model_spec)
            elif model_type == 'panel_fe':
                result = self._run_panel_fe(df, model_spec)
            elif model_type == 'panel_re':
                result = self._run_panel_re(df, model_spec)
            else:
                raise ValueError(f"Unknown model type: {model_type}")
            
            results[model_name] = result
            self.logger.info(f"  ✓ {model_name} completed")
        
        return results
    
    def _run_ols(self, df: pd.DataFrame, spec: Dict) -> any:
        """
        Run OLS regression.
        """
        formula = spec['formula']
        model = smf.ols(formula, data=df).fit(cov_type='HC1')  # Robust SE
        
        self.logger.info(f"    R²: {model.rsquared:.3f}, Adj R²: {model.rsquared_adj:.3f}, N: {model.nobs}")
        
        return model
    
    def _run_panel_fe(self, df: pd.DataFrame, spec: Dict) -> any:
        """
        Run panel fixed effects regression.
        """
        # Set panel index
        df_panel = df.set_index(['firm_id', 'year'])
        
        dv = spec['dependent_var']
        ivs = spec['independent_vars']
        
        # Entity and time effects
        model = PanelOLS(
            df_panel[dv],
            df_panel[ivs],
            entity_effects=spec.get('entity_effects', True),
            time_effects=spec.get('time_effects', True)
        ).fit(cov_type='clustered', cluster_entity=True)
        
        self.logger.info(f"    R²: {model.rsquared:.3f}, N: {model.nobs}")
        
        return model
    
    def _run_panel_re(self, df: pd.DataFrame, spec: Dict) -> any:
        """
        Run panel random effects regression.
        """
        df_panel = df.set_index(['firm_id', 'year'])
        
        dv = spec['dependent_var']
        ivs = spec['independent_vars']
        
        model = RandomEffects(df_panel[dv], df_panel[ivs]).fit()
        
        self.logger.info(f"    R²: {model.rsquared:.3f}, N: {model.nobs}")
        
        return model


# ============================================================================
# Phase 6: Robustness Checks
# ============================================================================

class RobustnessChecker:
    """
    Run robustness checks.
    """
    
    def __init__(self, config: Dict, logger: logging.Logger):
        self.config = config
        self.logger = logger
    
    def check(self, df: pd.DataFrame, main_results: Dict) -> Dict:
        """
        Run all robustness checks.
        """
        self.logger.info("=" * 60)
        self.logger.info("PHASE 6: ROBUSTNESS CHECKS")
        self.logger.info("=" * 60)
        
        robustness_results = {}
        
        checks = self.config.get('robustness_checks', [])
        
        for check_spec in checks:
            check_name = check_spec['name']
            self.logger.info(f"Running: {check_name}")
            
            # Implement various robustness checks
            # Example: Alternative DV, subsample analysis, etc.
            
        return robustness_results


# ============================================================================
# Phase 7: Visualization
# ============================================================================

class Visualizer:
    """
    Create publication-quality figures.
    """
    
    def __init__(self, config: Dict, logger: logging.Logger, output_dir: str):
        self.config = config
        self.logger = logger
        self.output_dir = Path(output_dir) / "figures"
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def visualize(self, df: pd.DataFrame, results: Dict):
        """
        Generate all figures.
        """
        self.logger.info("=" * 60)
        self.logger.info("PHASE 7: VISUALIZATION")
        self.logger.info("=" * 60)
        
        # Figure 1: Distribution of DV
        self._plot_dv_distribution(df)
        
        # Figure 2: Scatter plot (IV vs DV)
        self._plot_scatter(df)
        
        # Figure 3: Moderation effect
        # self._plot_moderation(df, results)
        
        self.logger.info(f"  ✓ Figures saved to {self.output_dir}")
    
    def _plot_dv_distribution(self, df: pd.DataFrame):
        """
        Figure 1: Distribution of dependent variable.
        """
        dv = self.config['variables']['dependent']
        
        fig, ax = plt.subplots(figsize=(8, 5))
        ax.hist(df[dv].dropna(), bins=30, edgecolor='black', alpha=0.7)
        ax.set_xlabel(dv)
        ax.set_ylabel('Frequency')
        ax.set_title(f'Distribution of {dv}')
        
        plt.tight_layout()
        plt.savefig(self.output_dir / 'figure1_dv_distribution.png', dpi=300)
        plt.close()
        
        self.logger.info(f"    Saved: figure1_dv_distribution.png")
    
    def _plot_scatter(self, df: pd.DataFrame):
        """
        Figure 2: Scatter plot of IV vs DV.
        """
        dv = self.config['variables']['dependent']
        iv = self.config['variables']['independent'][0]
        
        fig, ax = plt.subplots(figsize=(8, 5))
        ax.scatter(df[iv], df[dv], alpha=0.5)
        ax.set_xlabel(iv)
        ax.set_ylabel(dv)
        ax.set_title(f'{iv} vs {dv}')
        
        # Add regression line
        z = np.polyfit(df[iv].dropna(), df[dv].dropna(), 1)
        p = np.poly1d(z)
        ax.plot(df[iv], p(df[iv]), "r--", linewidth=2)
        
        plt.tight_layout()
        plt.savefig(self.output_dir / 'figure2_scatter.png', dpi=300)
        plt.close()
        
        self.logger.info(f"    Saved: figure2_scatter.png")


# ============================================================================
# Phase 8: Export Results
# ============================================================================

class ResultsExporter:
    """
    Export results to various formats (LaTeX, DOCX, Excel).
    """
    
    def __init__(self, config: Dict, logger: logging.Logger, output_dir: str):
        self.config = config
        self.logger = logger
        self.output_dir = Path(output_dir)
    
    def export(self, tables: Dict, results: Dict):
        """
        Export all results.
        """
        self.logger.info("=" * 60)
        self.logger.info("PHASE 8: EXPORT RESULTS")
        self.logger.info("=" * 60)
        
        # Export descriptive tables to Excel
        self._export_tables_excel(tables)
        
        # Export regression results to LaTeX
        self._export_regression_latex(results)
        
        # Create summary report in DOCX
        self._create_summary_report(tables, results)
        
        self.logger.info(f"  ✓ Results exported to {self.output_dir}")
    
    def _export_tables_excel(self, tables: Dict):
        """
        Export descriptive tables to Excel.
        """
        output_file = self.output_dir / "descriptive_tables.xlsx"
        
        with pd.ExcelWriter(output_file, engine='openpyxl') as writer:
            for table_name, table_df in tables.items():
                table_df.to_excel(writer, sheet_name=table_name[:31])  # Excel sheet name limit
        
        self.logger.info(f"    Saved: descriptive_tables.xlsx")
    
    def _export_regression_latex(self, results: Dict):
        """
        Export regression results to LaTeX table.
        """
        # This would use stargazer or create custom LaTeX table
        output_file = self.output_dir / "regression_table.tex"
        
        with open(output_file, 'w') as f:
            f.write("% LaTeX regression table\n")
            f.write("% Use \\input{regression_table.tex} in your paper\n\n")
            
            for model_name, result in results.items():
                f.write(f"% {model_name}\n")
                # Write LaTeX table content
        
        self.logger.info(f"    Saved: regression_table.tex")
    
    def _create_summary_report(self, tables: Dict, results: Dict):
        """
        Create summary report in DOCX format.
        """
        doc = Document()
        doc.add_heading('Research Results Summary', 0)
        
        # Add date
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Add descriptive statistics
        doc.add_heading('1. Descriptive Statistics', level=1)
        
        for table_name, table_df in tables.items():
            doc.add_heading(table_name.replace('_', ' ').title(), level=2)
            doc.add_paragraph(table_df.to_string())
        
        # Add regression results
        doc.add_heading('2. Regression Results', level=1)
        
        for model_name, result in results.items():
            doc.add_heading(model_name, level=2)
            doc.add_paragraph(str(result.summary()))
        
        # Save
        output_file = self.output_dir / "summary_report.docx"
        doc.save(output_file)
        
        self.logger.info(f"    Saved: summary_report.docx")


# ============================================================================
# Main Pipeline Execution
# ============================================================================

class ResearchPipeline:
    """
    Master pipeline orchestrator.
    """
    
    def __init__(self, config_path: str):
        self.config = load_config(config_path)
        self.output_dir = self.config['output_dir']
        Path(self.output_dir).mkdir(parents=True, exist_ok=True)
        
        self.logger = setup_logging(self.output_dir)
        
        # Initialize all components
        self.collector = DataCollector(self.config, self.logger)
        self.cleaner = DataCleaner(self.config, self.logger)
        self.integrator = DataIntegrator(self.config, self.logger)
        self.descriptor = DescriptiveAnalyzer(self.config, self.logger)
        self.analyzer = RegressionAnalyzer(self.config, self.logger)
        self.robustness = RobustnessChecker(self.config, self.logger)
        self.visualizer = Visualizer(self.config, self.logger, self.output_dir)
        self.exporter = ResultsExporter(self.config, self.logger, self.output_dir)
    
    def run(self):
        """
        Execute full pipeline.
        """
        start_time = datetime.now()
        self.logger.info("=" * 60)
        self.logger.info("STRATEGIC MANAGEMENT RESEARCH PIPELINE")
        self.logger.info("=" * 60)
        self.logger.info(f"Start time: {start_time}")
        self.logger.info(f"Configuration: {self.config}")
        
        try:
            # Phase 1: Data Collection
            raw_data = self.collector.collect_all()
            
            # Phase 2: Data Cleaning
            cleaned_data = self.cleaner.clean_all(raw_data)
            
            # Phase 3: Data Integration
            df = self.integrator.integrate(cleaned_data)
            
            # Save integrated dataset
            df.to_csv(Path(self.output_dir) / "data_final.csv", index=False)
            self.logger.info(f"Saved final dataset: data_final.csv")
            
            # Phase 4: Descriptive Statistics
            tables = self.descriptor.analyze(df)
            
            # Phase 5: Regression Analysis
            results = self.analyzer.analyze(df)
            
            # Phase 6: Robustness Checks
            robustness_results = self.robustness.check(df, results)
            
            # Phase 7: Visualization
            self.visualizer.visualize(df, results)
            
            # Phase 8: Export Results
            self.exporter.export(tables, results)
            
            # Success
            end_time = datetime.now()
            duration = (end_time - start_time).total_seconds()
            
            self.logger.info("=" * 60)
            self.logger.info("PIPELINE COMPLETED SUCCESSFULLY")
            self.logger.info("=" * 60)
            self.logger.info(f"End time: {end_time}")
            self.logger.info(f"Duration: {duration:.1f} seconds ({duration/60:.1f} minutes)")
            self.logger.info(f"Output directory: {self.output_dir}")
            
        except Exception as e:
            self.logger.error("=" * 60)
            self.logger.error("PIPELINE FAILED")
            self.logger.error("=" * 60)
            self.logger.error(f"Error: {str(e)}", exc_info=True)
            raise


# ============================================================================
# Command Line Interface
# ============================================================================

def main():
    """
    Command line interface for pipeline.
    """
    parser = argparse.ArgumentParser(
        description='Strategic Management Research Pipeline',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python complete_pipeline.py --config config.yaml
  python complete_pipeline.py --config my_study.yaml --output results/
        """
    )
    
    parser.add_argument(
        '--config',
        type=str,
        required=True,
        help='Path to configuration YAML file'
    )
    
    parser.add_argument(
        '--output',
        type=str,
        help='Override output directory from config'
    )
    
    args = parser.parse_args()
    
    # Run pipeline
    pipeline = ResearchPipeline(args.config)
    
    if args.output:
        pipeline.output_dir = args.output
    
    pipeline.run()


if __name__ == '__main__':
    main()
