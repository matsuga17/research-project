#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
最終論文作成スクリプト
既存のDraft1とDraft2+Draft3を統合し、Word文書として出力
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_final_paper():
    """最終論文をWord文書として作成"""
    
    # 既存のMarkdownファイルを読み込み
    draft1_path = "調査と分析/Week3_Day25-27asia_pacific_ecommerce_research_draft1.md"
    draft2_path = "調査と分析/Week3_Day28-29asia_pacific_ecommerce_research_draft2+draft3.md"
    
    print(f"Draft1を読み込み中: {draft1_path}")
    with open(draft1_path, 'r', encoding='utf-8') as f:
        draft1_content = f.read()
    
    print(f"Draft2+Draft3を読み込み中: {draft2_path}")
    with open(draft2_path, 'r', encoding='utf-8') as f:
        draft2_content = f.read()
    
    # Word文書を作成
    doc = Document()
    
    # 文書のプロパティを設定
    core_props = doc.core_properties
    core_props.author = "Changu"
    core_props.title = "アジア太平洋E-commerce市場における事業戦略の構造的差異"
    core_props.comments = "市場成熟度と戦略パターン適合性の実証的研究"
    
    # スタイルを設定
    styles = doc.styles
    
    # 通常スタイル
    style_normal = styles['Normal']
    style_normal.font.name = 'MS Mincho'
    style_normal.font.size = Pt(10.5)
    style_normal.paragraph_format.line_spacing = 1.5
    style_normal.paragraph_format.first_line_indent = Inches(0.17)  # 1文字分のインデント
    
    # タイトルページを作成
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("アジア太平洋E-commerce市場における\n事業戦略の構造的差異")
    title_run.font.name = 'MS Gothic'
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run("—市場成熟度と戦略パターン適合性の実証的研究—")
    subtitle_run.font.name = 'MS Mincho'
    subtitle_run.font.size = Pt(12)
    
    # 空行
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 著者情報
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_run = author_para.add_run("研究者：Changu\n作成日：2025年11月20日\n論文種別：探索的実証研究\n対象領域：国際経営戦略、プラットフォーム経済、デジタル・トランスフォーメーション")
    author_run.font.name = 'MS Mincho'
    author_run.font.size = Pt(10.5)
    
    # 改ページ
    doc.add_page_break()
    
    # Markdownコンテンツを簡易的に変換してWord文書に追加
    # （完全なMarkdown→Wordコンバータではなく、基本的な構造のみ処理）
    
    combined_content = draft1_content + "\n\n" + draft2_content
    
    for line in combined_content.split('\n'):
        line = line.strip()
        if not line:
            continue
        
        # 見出しレベルを判定
        if line.startswith('## '):
            # H2 - 大見出し
            para = doc.add_heading(line[3:], level=1)
            para.style.font.name = 'MS Gothic'
        elif line.startswith('### '):
            # H3 - 中見出し
            para = doc.add_heading(line[4:], level=2)
            para.style.font.name = 'MS Gothic'
        elif line.startswith('#### '):
            # H4 - 小見出し
            para = doc.add_heading(line[5:], level=3)
            para.style.font.name = 'MS Gothic'
        elif line.startswith('**') and line.endswith('**'):
            # 太字
            para = doc.add_paragraph()
            run = para.add_run(line[2:-2])
            run.bold = True
        elif line.startswith('　'):
            # 本文（インデントあり）
            doc.add_paragraph(line)
        else:
            # その他のテキスト
            doc.add_paragraph(line)
    
    # 保存
    output_path = "final_paper.docx"
    doc.save(output_path)
    print(f"\n最終論文を作成しました: {output_path}")
    print(f"ファイルサイズ: {os.path.getsize(output_path) / 1024:.2f} KB")
    
    return output_path

if __name__ == "__main__":
    try:
        create_final_paper()
    except Exception as e:
        print(f"エラーが発生しました: {e}")
        import traceback
        traceback.print_exc()
